import docx
import re

# TODO: execption handle

version = '0.1'

chip_name = 'mc32p7510'
doc_file = r'E:\temp\jupyter\MC32P7510_UMAN_V1.5.docx'

REG_SUMMARY_TABLE_IDENTIFY = '通用数据区'
IGNORE_REG_NAME = ['', '通用数据区', '保留']


class WordTable():

    def __init__(self, table):
        self.table = table

    def cell_text(self, idx_row, idx_col):
        return self.table.cell(idx_row, idx_col).text

    def has_text(self, str):
        for i in range(len(self.table.rows)):
            for j in range(len(self.table.columns)):
                if str == self.cell_text(i, j):
                    return True
        return False


class RegSummaryTable(WordTable):

    def __found_addr_range(self, text):
        # match like '188H - 18FH'
        if re.match('[0-9ABCDEF]+H – [0-9ABCDEF]+H', text):
            return True
        else:
            return False

    def __is_reg_list_row(self, idx_row):
        row_cells = self.table.row_cells(idx_row)
        first_cell_txt = row_cells[0].text
        if len(row_cells) == 9 and self.__found_addr_range(first_cell_txt):
            return True
        else:
            return False

    def __get_start_addr(self, idx_row):
        # row_cells = self.table.row_cells(idx_row)
        first_cell_txt = self.table.row_cells(idx_row)[0].text
        m = re.match('([0-9ABCDEF]+)H – [0-9ABCDEF]+H', first_cell_txt)
        start_addr_txt = m.group(1) if m else '0'
        return int('0x' + start_addr_txt, 16)

    def get_reg_list_row(self, idx_row):
        reg_lst = []
        row_cells = self.table.row_cells(idx_row)
        start_addr = self.__get_start_addr(idx_row)
        for i in range(1, len(row_cells)):
            if (row_cells[i].text not in IGNORE_REG_NAME):
                reg_lst.append((start_addr + i - 1, row_cells[i].text))
        return reg_lst

    def get_reg_list_all(self):
        all_reg_list = []
        for i_row in range(len(self.table.rows)):
            if self.__is_reg_list_row(i_row):
                all_reg_list.append(self.get_reg_list_row(i_row))
        return all_reg_list


def search_reg_summary_table(doc):
    for tab in doc.tables:
        wt = WordTable(tab)
        if wt.has_text(REG_SUMMARY_TABLE_IDENTIFY):
            return tab
    return None


def search_reg_bit_table(doc, reg_name):
    for tab in doc.tables:
        wt = WordTable(tab)
        if wt.cell_text(1, 0).strip() == reg_name:
            return tab
    return None


def get_a_reg_bit_def(tab):
    bit_def_list = []
    for i in range(8):
        bit_def_list.append((7 - i, tab.cell(1, i + 1).text))
    return bit_def_list


def get_all_reg_info(doc, all_reg_list):
    all_reg_info = []
    for reg_lst in all_reg_list:
        for (addr, reg_name) in reg_lst:
            a_reg_info = {}
            a_reg_info['addr'] = addr
            a_reg_info['reg_name'] = reg_name
            reg_bit_tab = search_reg_bit_table(doc, reg_name)
            if reg_bit_tab:
                a_reg_info['bit_def'] = get_a_reg_bit_def(reg_bit_tab)
            else:
                a_reg_info['bit_def'] = None
            all_reg_info.append(a_reg_info)
    return all_reg_info


def write_header(f):
    f.write('#include "reg_def.h"\n\n')
    f.write(
        '#define SET_BITS(x, n, of) (~((~(((-1) << (n)) | (x))) << (of))) & 0xFFFF\n\n')
    f.write('const struct reg_def all_reg[] = {\n')


def write_tail(f):
    f.write('};\n\n')
    f.write('const int num_reg = sizeof(all_reg) / sizeof(all_reg[0]);\n\n')
    f.write('const int maxram = 0x0ff;\n')
    f.write(
        'const int badram[][2] = {{0x80, 0x81}, {0x85, 0x85}, {-1, -1}};\n\n')
    f.write('const struct bit_def config1[] = {};\n')
    f.write('const int num_config1 = sizeof(config1) / sizeof(config1[0]);\n')
    f.write('const int config1_addr = 0x8001;\n\n')
    f.write('const struct bit_def config0[] = {};\n')
    f.write('const int num_config0 = sizeof(config0) / sizeof(config0[0]);\n')
    f.write('const int config0_addr = 0x8000;\n')


def write_bit(f, bit_name):
    if bit_name.strip() == '' or bit_name.strip() == '-':
        f.write('{"", 0}')
    else:
        f.write('{"%s", 1}' % bit_name)


def write_reg_line_with_bit(f, a_reg_info):
    f.write('\t{"%s", 0x%04x, 1, {' %
            (a_reg_info['reg_name'], a_reg_info['addr']))
    for i in range(7):
        bitn, bit_name = a_reg_info['bit_def'][7 - i]
        write_bit(f, bit_name)
        f.write(', ')
    bitn, bit_name = a_reg_info['bit_def'][0]
    write_bit(f, bit_name)
    f.write('}}')


def write_reg_line(f, a_reg_info):
    if a_reg_info['bit_def'] is None:
        f.write('\t{"%s", 0x%04x, 0, {{"", 0}, {"", 0}, {"", 0}, {"", 0}, {"", 0}, {"", 0}, {"", 0}, {"", 0}}}' %
                (a_reg_info['reg_name'], a_reg_info['addr']))
    else:
        write_reg_line_with_bit(f, a_reg_info)


def write_reg_file(chip_name, all_reg_info):
    out_file = 'reg_def_%s.c' % chip_name
    with open(out_file, 'w+') as f:
        write_header(f)
        for i in range(len(all_reg_info) - 1):
            write_reg_line(f, all_reg_info[i])
            f.write(',\n')
        write_reg_line(f, all_reg_info[-1])
        f.write('\n')
        write_tail(f)


def main():
    doc = docx.Document(doc_file)
    reg_sum_tab = RegSummaryTable(search_reg_summary_table(doc))
    if reg_sum_tab is None:
        print('Register summary table not found!')
        return -1
    else:
        all_reg_list = reg_sum_tab.get_reg_list_all()
        all_reg_info = get_all_reg_info(doc, all_reg_list)
        # print(all_reg_info)
        write_reg_file(chip_name, all_reg_info)
        print('done!')
        return 0


if __name__ == '__main__':
    main()
